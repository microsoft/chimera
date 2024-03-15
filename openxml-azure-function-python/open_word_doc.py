import re
import io
from azure.storage.blob import BlobServiceClient
from docx import Document
from collections import defaultdict

TAK_REGEX_PATTERN = r"TAK-(\d+)"
TKD_REGEX_PATTERN = r"TKD-\w+-(\d+)"

def create_blob_client(connection_string, container_name, blob_name):
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    blob_client = blob_service_client.get_blob_client(container_name, blob_name)
    return blob_client

def read_word_document_from_blob_storage(connection_string, container_name, blob_name):
    blob_client = create_blob_client(connection_string, container_name, blob_name)

    download_stream = blob_client.download_blob().readall()
    document = Document(io.BytesIO(download_stream))

    values = split_sections(document)
    headers = pull_headers(document)
    return values, headers

def pull_headers(document):
    headers = defaultdict(str)
    reg_tak = re.compile(TAK_REGEX_PATTERN)
    reg_tkd = re.compile(TKD_REGEX_PATTERN)

    for paragraph in document.paragraphs:
        add_match_to_headers(headers, "TAK", reg_tak.search(paragraph.text))
        add_match_to_headers(headers, "TKD", reg_tkd.search(paragraph.text))

    return headers

def add_match_to_headers(headers, key, match):
    if match:
        headers[key] = match.group()

def split_sections(document):
    sections = defaultdict(str)
    current_key = None
    paragraph_text = ""

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            if current_key:
                sections[current_key] = paragraph_text.rstrip('|')
                paragraph_text = ""
            current_key = paragraph.text
        elif current_key and paragraph.text:
            paragraph_text += paragraph.text + "||"

    if current_key:
        sections[current_key] = paragraph_text.rstrip('|')

    return sections
