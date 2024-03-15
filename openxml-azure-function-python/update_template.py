import re
import io
from azure.storage.blob import BlobServiceClient
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from collections import defaultdict

def create_blob_client(connection_string, container_name, blob_name):
    blob_service_client = BlobServiceClient.from_connection_string(connection_string)
    blob_client = blob_service_client.get_blob_client(container_name, blob_name)
    return blob_client

def update_document_template(sections, headers, connection_string, source_container_name, source_blob_name, destination_container_name, destination_blob_name):
    source_blob_client = create_blob_client(connection_string, source_container_name, source_blob_name)

    download_stream = source_blob_client.download_blob().readall()
    document = Document(io.BytesIO(download_stream))

    # Replace headers
    replace_header(document, headers)

    # Replace sections
    for key, value in sections.items():
        key = key.upper()
        if key == "INTRODUCTION":
            replace_sections(document, "##INTRODUCTION##", value.replace("||", ""))
        elif key == "STUDY OBJECTIVE":
            replace_sections(document, "##OBJECTIVES##", value.replace("||", ""))
        elif key == "MATERIALS":
            replace_sections(document, "##MATERIALS##", value.replace("||", ""))
        elif key == "TEST SYSTEM":
            replace_sections(document, "##TESTSYSTEM##", value.replace("||", ""))
        elif key == "TESTING FACILITY":
            replace_sections(document, "##TESTING FACILITY##", value.replace("||", ""))
        elif key == "TEST ARTICLE":
            replace_sections(document, "##Test Article##", value.replace("||", ""))
        elif key == "ARCHIVING":
            replace_sections(document, "##ARCHIVING##", value.replace("||", ""))
        elif key == "GOOD LABORATORY PRACTICE COMPLIANCE":
            glp = ""
            if "this study was not conducted in accordance" in value.lower() or "this study will not be conducted in accordance" in value.lower():
                glp = "This study was not conducted in accordance with the Code of Federal Regulations (CFR), Title 21, Part 58: Good Laboratory Practice (GLP) for Nonclinical Laboratory Studies, issued by the United States Food and Drug Administration (FDA). This study was conducted as a basic exploratory study and as such, the data from this study were not audited by Quality Assurance (QA). However, all data were recorded appropriately and documentation necessary for reconstruction of this study is available in the study files at Takeda Development Center Americas, Inc. (TDCA) (San Diego, CA, USA). The data are accurately reflected in the report."
            else:
                glp = "This study was conducted in accordance with the Code of Federal Regulations, Title 21, Part 58: Good Laboratory Practice for Nonclinical Laboratory Studies, issued by the United States Food and Drug Administration. The study was conducted to meet GLP standards, and all data were audited by Quality Assurance to ensure accuracy and compliance. The necessary documentation for reconstruction of this study is available in the study files at Takeda Development Center Americas, Inc. (TDCA) (CITY, ST, USA). The data presented in this report accurately reflect the findings of the study."
            replace_sections(document, "##GOODLABCOMPLIANCE##", glp)
        elif key == "AMENDMENTS TO, AND DEVIATIONS FROM, THE PROTOCOL":
            replace_sections(document, "##DEVIATIONSFROMPROTOCOL##", value.replace("||", ""))
        elif key == "SUMMARY":
            replace_sections(document, "##AISUMMARY##", value.replace("||", ""))
        else:
            # Non standard headers to be evaluated here.
            if key.startswith("DOCUMENT TITLE"):
                document_title_replace(document, value)
            #This generates an error
            # if key.startswith("PROTOCOL APPROVAL"):
            #     # need to split the section value into parts
            #     parts = value.split("||")
            #     add_approvals(document, parts)

    # Save the updated document to a new blob
    destination_blob_client = create_blob_client(connection_string, destination_container_name, destination_blob_name)
    with io.BytesIO() as buf:
        document.save(buf)
        buf.seek(0)
        destination_blob_client.upload_blob(buf, overwrite=True)

def document_title_replace(document, section):
    title_page = section[1].split("||")
    replace_sections(document, "TITLE:##TITLE##", "TITLE: " + section[0][16:])
    replace_sections(document, "##TITLE##", section[0][16:])
    study_dir = next((s for s in title_page if "Study Director" in s), None)
    if study_dir is not None:
        split_value = study_dir.split(":")
        if len(split_value) > 1:
            find_and_replace(document, "STUDYDIRECTOR", split_value[1])

def replace_sections(document, search_text, update_text):
    for paragraph in document.paragraphs:
        if search_text in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(search_text, update_text)

def find_and_replace(document, search_text, replace_value):
    for paragraph in document.paragraphs:
        if search_text in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(search_text, replace_value)

def add_approvals(document, update_text_list):
    old_table = next((table for table in document.tables if "##APPROVED BY##" in table.text), None)
    if old_table is not None:
        new_table = document.add_table(rows=0, cols=1)
        for i in range(0, len(update_text_list), 4):
            create_new_row(new_table, update_text_list[i], True)
            create_new_row(new_table, update_text_list[i + 2])
            create_new_row(new_table, update_text_list[i + 3])
            create_new_row(new_table, " ")
        document.element.replace(old_table._element, new_table._element)

def create_new_row(table, text, highlighted=False):
    row = table.add_row()
    cell = row.cells[0]
    run = cell.paragraphs[0].add_run(text)
    if highlighted:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        run.font.size = Pt(22)  # Font size
        run.italic = True  # Italic text
        # Add bottom border to the cell
        border_xml = '<w:bottom w:val="single" w:sz="12" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" />'
        cell._element.get_or_add_tcPr().append(parse_xml(border_xml))

def replace_header(document, header_values):
    # Get the header parts
    header_parts = document.sections[0].header

    # Loop through each header part
    for paragraph in header_parts.paragraphs:
        for run in paragraph.runs:
            if "TAK-" in run.text:
                run.text = run.text.replace("TAK-", header_values["TAK"])
            if "XXX" in run.text:
                run.text = run.text.replace("XXX", "")
            if "TKD-BCS-" in run.text:
                run.text = run.text.replace("TKD-BCS-", header_values["TKD"])
            if "XX-RX" in run.text:  # Three X's have already been removed
                run.text = run.text.replace("XX-RX", "")
